from __future__ import annotations

import math
import os
import random
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = ROOT / "tmp" / "doc_label_previews"
DOCX_PATH = OUT_DIR / "Punit_Label_Application_Flow_and_Label_Format_Guide.docx"

INK = (32, 32, 32)
MUTED = (85, 85, 85)
PAPER = (253, 253, 251)
BORDER = (218, 220, 224)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def fit_font(text: str, max_width: int, preferred: int, minimum: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    size = preferred
    probe = Image.new("RGB", (10, 10))
    draw = ImageDraw.Draw(probe)
    while size > minimum:
        f = font(size, bold)
        if draw.textbbox((0, 0), text, font=f)[2] <= max_width:
            return f
        size -= 1
    return font(minimum, bold)


def draw_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, size: int, bold: bool = False, fill=INK) -> None:
    draw.text(xy, text, font=font(size, bold), fill=fill)


def draw_barcode(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, h: int, data: str) -> None:
    draw.rectangle((x, y, x + w, y + h), fill=(255, 255, 255))
    seed = sum((idx + 1) * ord(ch) for idx, ch in enumerate(data))
    rng = random.Random(seed)
    cursor = x + 8
    while cursor < x + w - 8:
        bar = rng.choice([2, 3, 4, 5, 7])
        gap = rng.choice([2, 3, 4])
        draw.rectangle((cursor, y + 2, min(cursor + bar, x + w - 8), y + h - 2), fill=(0, 0, 0))
        cursor += bar + gap


def bordered_canvas(width: int, height: int) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), PAPER)
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((2, 2, width - 3, height - 3), radius=18, outline=(170, 170, 170), width=2)
    return img, draw


def generic_preview(
    name: str,
    width: int,
    height: int,
    attrs: list[tuple[str, str]],
    *,
    grid: bool,
    business_hours: str = "",
    white_label: bool = False,
    print_time: bool = True,
) -> Path:
    img, draw = bordered_canvas(width, height)
    y = 22
    if not white_label:
        draw_text(draw, (22, y), "Punit Foods Pvt. Ltd.", 36, True)
        y += 54
        for line in ["ops@punitfoods.in | +91 98765 43210", "GST: 24AABCP1234Z1ZV"]:
            draw_text(draw, (22, y), line, 22, False, MUTED)
            y += 28
        if business_hours:
            draw_text(draw, (22, y), business_hours, 22, False, MUTED)
            y += 34
    if print_time:
        draw_text(draw, (width - 160, 24), "22-06-2026", 18)
        draw_text(draw, (width - 160, 46), "14:35", 18)

    if business_hours:
        title = "Wholesale Pack"
        title_font = font(48, True)
        tw = draw.textbbox((0, 0), title, font=title_font)[2]
        draw.text(((width - tw) / 2, y), title, font=title_font, fill=INK)
        y += 72
        product_size = 46
        line_height = 60
        col_gap = 330
        bottom_padding = 185
        barcode_height = 92
    elif width == 600 and height == 410:
        product_size = 31
        line_height = 40
        col_gap = 140
        bottom_padding = 80
        barcode_height = 46
    else:
        product_size = 32
        line_height = 56
        col_gap = 190
        bottom_padding = 150
        barcode_height = 80

    product = name
    pf = fit_font(product, width - 44, product_size, 22, True)
    draw.text((22, y), product, font=pf, fill=INK)
    y += 46 if not business_hours else 72

    if not grid:
        for key, value in attrs[:10]:
            if y > height - bottom_padding:
                break
            draw_text(draw, (22, y), f"{key}:", 24 if width == 600 else 28)
            draw_text(draw, (22 + col_gap, y), value, 26 if width == 600 else 30, True)
            y += line_height
    else:
        x1 = 22
        x2 = width // 2 + 20
        for i in range(0, min(len(attrs), 10), 2):
            if y > height - bottom_padding:
                break
            first = attrs[i]
            draw_text(draw, (x1, y), f"{first[0]}:", 22)
            draw_text(draw, (x1 + col_gap, y), first[1], 23, True)
            if i + 1 < len(attrs):
                second = attrs[i + 1]
                draw_text(draw, (x2, y), f"{second[0]}:", 22)
                draw_text(draw, (x2 + col_gap, y), second[1], 23, True)
            y += line_height

    footer_y = height - 34
    barcode_y = footer_y - barcode_height - 18
    draw_barcode(draw, 22, barcode_y, width - 44, barcode_height, "A7K91")
    if not white_label and not business_hours:
        footer = "Label generated from weighing ERP by punitinstrument.com"
        ff = font(17)
        tw = draw.textbbox((0, 0), footer, font=ff)[2]
        draw.text(((width - tw) / 2, footer_y), footer, font=ff, fill=MUTED)

    path = ASSET_DIR / f"{name.lower().replace(' ', '_')}.png"
    img.save(path)
    return path


def tea_preview() -> Path:
    img, draw = bordered_canvas(600, 410)
    draw_text(draw, (35, 35), "Majedar Tea Co.", 40, True)
    draw_text(draw, (370, 35), "22-06-2026 14:35", 21)
    draw_text(draw, (25, 100), "Premium Assam Tea", 32, True)
    draw_text(draw, (25, 150), "Grade : BOP", 30, True)
    draw_text(draw, (25, 196), "Weight :", 30, True)
    draw_text(draw, (260, 196), "1.250", 30, True)
    draw_barcode(draw, 50, 306, 500, 52, "A7K91")
    path = ASSET_DIR / "majedar_tea.png"
    img.save(path)
    return path


def dry_fruit_preview() -> Path:
    img, draw = bordered_canvas(600, 410)
    draw_text(draw, (345, 15), "22-06-2026 14:35", 20, True)
    draw_text(draw, (25, 46), "Premium Cashew W320", 38, True)
    rows = [("Grade", "W320"), ("Batch", "DF-2406"), ("Net Weight", "1.250 Kg"), ("Units", "25")]
    y = 96
    for key, value in rows:
        draw_text(draw, (25, y), f"{key} : {value}", 31, True)
        y += 42
    draw_barcode(draw, 50, 280, 500, 80, "A7K91")
    path = ASSET_DIR / "dry_fruit.png"
    img.save(path)
    return path


def small_seven_preview() -> Path:
    img, draw = bordered_canvas(600, 410)
    draw_text(draw, (30, 20), "Punit Foods Pvt. Ltd.", 34, True)
    draw_text(draw, (435, 72), "22-06-2026", 18)
    draw_text(draw, (435, 96), "14:35", 18)
    draw_text(draw, (30, 72), "Spice Mix 7", 28, True)
    draw_text(draw, (30, 112), "Sr No : 1", 19)
    left_x, right_x, y = 30, 310, 146
    rows = [("Color", "Red"), ("Size", "Small"), ("Units", "25"), ("Net Weight", "1.250")]
    for idx in range(0, len(rows), 2):
        a = rows[idx]
        draw_text(draw, (left_x, y), f"{a[0]} : {a[1]}", 24)
        if idx + 1 < len(rows):
            b = rows[idx + 1]
            draw_text(draw, (right_x, y), f"{b[0]} : {b[1]}", 24)
        y += 36
    draw_barcode(draw, 45, 282, 510, 46, "A7K91")
    path = ASSET_DIR / "small_seven.png"
    img.save(path)
    return path


def runtime_preview() -> Path:
    img, draw = bordered_canvas(600, 600)
    draw_text(draw, (28, 28), "Custom 75x75 Template", 31, True)
    draw_text(draw, (28, 82), "company_name: Punit Foods", 22)
    draw_text(draw, (28, 125), "product_name: Besan Flour", 27, True)
    draw_text(draw, (28, 176), "gross_weight: 1.500", 24)
    draw_text(draw, (28, 218), "tare_weight: 0.250", 24)
    draw_text(draw, (28, 260), "net_weight: 1.250", 24)
    draw_text(draw, (28, 312), "attr_grade: Grade: A", 23)
    draw_text(draw, (28, 350), "sr_no: 1", 23)
    draw_barcode(draw, 65, 412, 470, 90, "A7K91")
    draw_text(draw, (155, 530), "footer: Thank you", 22, False, MUTED)
    path = ASSET_DIR / "custom_runtime_template.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DADCE0", size="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_col_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_col_widths(t, widths)
    for idx, header in enumerate(headers):
        cell = t.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "F8F9FA")
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.name = "Arial"
            paragraph.runs[0].font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return t


def add_h(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0 if level < 3 else 67)
        run.font.bold = False


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_preview(doc: Document, title: str, path: Path, width: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    doc.add_picture(str(path), width=Inches(width))


def build_doc(previews: dict[str, Path]) -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.line_spacing = 1.15
    run = title.add_run("Punit Label Application Flow and Label Format Guide")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = False

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    srun = subtitle.add_run("Google Docs-ready guide generated from the Flutter source on 22 June 2026.")
    srun.font.name = "Arial"
    srun.font.size = Pt(11)
    srun.font.color.rgb = RGBColor(85, 85, 85)

    add_h(doc, "1. Application Flow", 1)
    for item in [
        "Splash and login: the app starts at SplashView, then routes to LoginView when credentials are required. Login collects Company Code, email, and password, then stores authenticated user and token data.",
        "Dashboard: DashBoardView loads user permissions, dashboard totals, top products, low-stock products, company details, device reconnection, and printer status polling. Role flags control whether Inward and Dispatch cards are accessible.",
        "Global drawer configuration: the drawer controls default non-batch label format, white-label mode, serial number printing, timestamp printing, label copy count, tare behavior, printer/receipt mode, tower light status, scale connection, and printer connection.",
        "Inward entry: operators select Batch Inward or Non Batch Wise Inward. Both flows capture gross, tare, and net weights, create barcodes from serial numbers, print labels, maintain logs, and submit pause/stop payloads to the API.",
        "Dispatch: operators select a customer, scan or manually enter a barcode, verify and add dispatch items, then save the dispatch transaction and export reports.",
    ]:
        add_number(doc, item)

    add_h(doc, "2. Configuration That Changes Labels", 1)
    config_rows = [
        ["White Label", "Removes company header/footer from generic and special previews when enabled."],
        ["Serial Number", "Prepends Sr No to labelFields where the selected label format supports it."],
        ["Time Stamp", "Adds date/time to static previews and print payloads."],
        ["Label Copies", "Repeats the platform print call from 1 to 10 times."],
        ["Tare State: OFF", "Tare is forced to 0 and labels usually print Net Weight only."],
        ["Tare State: ON", "Gross Weight, Tare Weight, and Net Weight can be printed."],
        ["Tare State: Barcode", "Tare products can be selected/scanned and included in weighing logic."],
        ["Label vs Receipt", "Label mode calls Android label-printer methods; Receipt mode sends item lines to the Bluetooth receipt printer."],
    ]
    table(doc, ["Setting", "Effect"], config_rows, [1.8, 4.7])

    add_h(doc, "3. Inward Workflow", 1)
    add_h(doc, "Batch Wise", 2)
    for item in [
        "Select a batch, then a product inside that batch.",
        "Select an existing label format. Batch flow currently previews existing/static formats.",
        "The controller syncs product tare unless global tare is OFF.",
        "Auto-weight mode monitors the scale every second. When the configured weight condition is met for the configured number of seconds, addToList creates a barcode, inserts a log entry, and prints the selected label.",
        "Pause/Stop groups barcodes by batch product and submits gross/tare/net/serial payloads, then generates an inward PDF report.",
    ]:
        add_bullet(doc, item)

    add_h(doc, "Non Batch Wise", 2)
    for item in [
        "Enter a transaction name, select product, select label format, and choose product attributes.",
        "When custom label templates are supported, the UI adds Custom, label size, and custom option selectors.",
        "The app groups repeated entries by same product and same selected attributes; serial numbers continue inside that group.",
        "Existing/static labels build labelFields from selected attributes plus weight fields. Custom templates resolve field positions and values from the runtime template payload.",
        "Pause/Stop submits the transaction payload and generates a transaction PDF report.",
    ]:
        add_bullet(doc, item)

    add_h(doc, "4. Dispatch Workflow", 1)
    for item in [
        "Open Dispatch from the dashboard when dispatch permission is enabled.",
        "Select a customer, then scan a barcode or type it into the Barcode field.",
        "verifyAndAddBarcode validates the barcode against dispatch data and appends it to dispatch logs.",
        "Save submits scanned barcodes. Export helpers can generate grouped PDF/Excel dispatch reports.",
    ]:
        add_bullet(doc, item)

    add_h(doc, "5. Label Format Catalog", 1)
    label_rows = [
        ["Majedar Tea", "600 x 410", "1", "First non-weight attribute becomes the description; net weight prints as Weight; Code128 barcode."],
        ["Small", "600 x 410", "3", "Generic label. Non-white mode allows up to 3 selectable fields; preview switches to grid when more than 1 field is present."],
        ["Medium", "600 x 600", "4 or 6", "Generic square label. Grid mode starts after 5 fields; white-label mode raises selectable fields to 6."],
        ["Large", "700 x 600", "5 or 8", "Generic larger label. Default non-batch fallback is Large. Grid mode starts after 5 fields."],
        ["Extra Large", "700 x 700", "7 or 9", "Generic large square label. Grid mode starts after 5 fields."],
        ["Wholesale Pack", "700 x 1200", "10", "Tall generic label with Wholesale Pack title, business-hours line, product name, selected attributes, Gross Weight, and barcode."],
        ["Small Seven", "600 x 410", "5", "Special two-column label. Excludes Weight and Sr No from attributes; unit conversion can insert Units."],
        ["DryFruit", "600 x 410", "3", "Special dry-fruit layout. Excludes gross/tare/weight/sr from attribute list; weight fields are formatted to 3 decimals with Kg."],
        ["Custom Runtime", "50x75, 75x75, 75x100, 100x100", "Template driven", "Fields come from API template coordinates: company, product, gross/tare/net, barcode, sr_no, datetime, footer, and attr_* fields."],
    ]
    table(doc, ["Format", "Preview/print size", "Selectable field limit", "Behavior"], label_rows, [1.25, 1.35, 1.35, 2.55])

    add_h(doc, "6. Label Field Rules", 1)
    for item in [
        "Small labels use Weight/Net Weight depending on the controller path; Medium, Large, and Extra Large usually include selected attributes plus Net Weight or Gross/Tare/Net depending on tare setting.",
        "Wholesale Pack keeps selected attributes and Gross Weight; Sr No is preserved when enabled.",
        "Majedar Tea ignores Sr No and weight keys when choosing its visible attribute description.",
        "DryFruit and SmallSeven remove weight-like keys from their attribute lists, then add their own weight/unit presentation.",
        "Runtime custom templates resolve field values by field_key. Supported built-in keys include company_name, company_email, company_contact_no, company_gst_no, company_website, company_address, product_name, gross_weight, tare_weight, net_weight, barcode, barcode_text, sr_no, datetime, footer, and attr_*.",
    ]:
        add_bullet(doc, item)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_h(doc, "7. Label Preview Examples", 1)
    add_para(doc, "The previews below use sample product data and recreate the app preview logic: label dimensions, visible fields, special filtering, grid thresholds, timestamp placement, and barcode placement.")

    add_preview(doc, "Majedar Tea Label Format", previews["tea"], 4.8)
    add_preview(doc, "Small Label Select Max (3)", previews["small"], 4.8)
    add_preview(doc, "Medium Label Select Max (4/6)", previews["medium"], 4.1)
    add_preview(doc, "Large Label Select Max (5/8)", previews["large"], 4.8)
    add_preview(doc, "Extra Large Label Select Max (7/9)", previews["extra_large"], 4.4)
    add_preview(doc, "Wholesale Pack", previews["wholesale"], 3.5)
    add_preview(doc, "Small Seven (5)", previews["small_seven"], 4.8)
    add_preview(doc, "DryFruit Label Format", previews["dry_fruit"], 4.8)
    add_preview(doc, "Custom Runtime Template Example", previews["runtime"], 4.4)

    add_h(doc, "8. Example Print Payload", 1)
    payload_rows = [
        ["Product", "Besan Flour"],
        ["Barcode", "A7K91"],
        ["Attributes", "Grade: A, Pack: 1 kg, Color: Yellow"],
        ["Gross Weight", "1.500"],
        ["Tare Weight", "0.250"],
        ["Net Weight", "1.250"],
        ["Units", "25, when unitConversion is enabled and unitValue x net weight equals 25"],
        ["Serial Number", "1, when Serial Number toggle is enabled"],
    ]
    table(doc, ["Field", "Example"], payload_rows, [1.65, 4.85])

    doc.save(DOCX_PATH)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    previews = {
        "tea": tea_preview(),
        "small": generic_preview(
            "Besan Flour",
            600,
            410,
            [("Weight", "1.250")],
            grid=False,
        ),
        "medium": generic_preview(
            "Rice Premium",
            600,
            600,
            [("Grade", "A"), ("Pack", "5 kg"), ("Net Weight", "5.000"), ("Batch", "B-2406")],
            grid=False,
        ),
        "large": generic_preview(
            "Wheat Flour",
            700,
            600,
            [("Grade", "A"), ("Pack", "10 kg"), ("Gross Weight", "10.250"), ("Tare Weight", "0.250"), ("Net Weight", "10.000")],
            grid=False,
        ),
        "extra_large": generic_preview(
            "Mixed Pulses",
            700,
            700,
            [("Grade", "A"), ("Pack", "10 kg"), ("Color", "Yellow"), ("Origin", "Gujarat"), ("Gross", "10.250"), ("Tare", "0.250"), ("Net", "10.000")],
            grid=True,
        ),
        "wholesale": generic_preview(
            "Wholesale Spice Box",
            700,
            1200,
            [("Sr No", "1"), ("Grade", "Export"), ("Pack", "25 kg"), ("Gross Weight", "25.800")],
            grid=False,
            business_hours="On working day 11:00AM - 6:00PM",
        ),
        "small_seven": small_seven_preview(),
        "dry_fruit": dry_fruit_preview(),
        "runtime": runtime_preview(),
    }
    build_doc(previews)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
